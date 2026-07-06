import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_margins(doc):
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

def add_header(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT, size=11, bold=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.bold = bold
    p.alignment = align
    
    # Add a bottom border-like effect using underline if needed, but simple bold is cleaner.
    return p

def add_paragraph(doc, text, size=10, bold=False, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    return p

def add_experience(doc, title, role_date, bullets):
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(0)
    p1.paragraph_format.space_before = Pt(4)
    run1 = p1.add_run(title)
    run1.font.name = 'Arial'
    run1.font.size = Pt(10)
    run1.bold = True
    
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    p2.paragraph_format.space_before = Pt(0)
    run2 = p2.add_run(role_date)
    run2.font.name = 'Arial'
    run2.font.size = Pt(10)
    run2.italic = True
    
    for bullet in bullets:
        add_bullet(doc, bullet)

def main():
    doc = docx.Document()
    set_margins(doc)
    
    # Contact Info
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run_name = p.add_run('Junyong Park\n')
    run_name.font.name = 'Arial'
    run_name.font.size = Pt(16)
    run_name.bold = True
    
    run_contact = p.add_run('London, United Kingdom | 6unyong@gmail.com | +44 7907303613 | linkedin.com/in/junyong-park-a00a3a228')
    run_contact.font.name = 'Arial'
    run_contact.font.size = Pt(10)
    
    # Summary
    add_header(doc, 'SUMMARY', size=12)
    add_paragraph(doc, 'Data Scientist and Data Analyst with practical experience in machine learning and deep learning. Skilled at finding business insights from complex data. Proven ability to solve real world problems using models like CNNs, Vision Language Models, and XGBoost. Strong focus on turning data analysis into practical business strategies.', space_after=6)
    
    # Core Skills
    add_header(doc, 'CORE SKILLS', size=12)
    add_bullet(doc, 'Languages and Tools: Python, SQL, Tableau, QGIS, AWS, Pandas, NumPy, Scikit Learn, PyTorch.')
    add_bullet(doc, 'Machine Learning and AI: Deep Learning, Computer Vision, Vision Language Models, Natural Language Processing, XGBoost, Clustering, Regression.')
    add_bullet(doc, 'Certifications: IELTS Academic, proDS Associate, Tableau Desktop Specialist, SQLD.')
    
    # Education
    add_header(doc, 'EDUCATION', size=12)
    
    p_edu1 = doc.add_paragraph()
    p_edu1.paragraph_format.space_after = Pt(0)
    p_edu1.paragraph_format.space_before = Pt(2)
    run_e1 = p_edu1.add_run('King’s College London, London, United Kingdom')
    run_e1.font.name = 'Arial'
    run_e1.font.size = Pt(10)
    run_e1.bold = True
    add_paragraph(doc, 'Master of Science in Data Science (Sep 2025 to Sep 2026)', space_after=4)
    
    p_edu2 = doc.add_paragraph()
    p_edu2.paragraph_format.space_after = Pt(0)
    p_edu2.paragraph_format.space_before = Pt(0)
    run_e2 = p_edu2.add_run('Kookmin University, Seoul, South Korea')
    run_e2.font.name = 'Arial'
    run_e2.font.size = Pt(10)
    run_e2.bold = True
    add_paragraph(doc, 'Bachelor of Business Administration (Mar 2019 to Jul 2025)', space_after=0)
    add_paragraph(doc, 'Major in Management Information System, Minor in Business Analytics and Statistics', space_after=6)
    
    # Patents
    add_header(doc, 'PATENTS', size=12)
    p_pat = doc.add_paragraph()
    p_pat.paragraph_format.space_after = Pt(6)
    p_pat.paragraph_format.space_before = Pt(2)
    run_pat = p_pat.add_run('Patent Application: ')
    run_pat.bold = True
    run_pat.font.name = 'Arial'
    run_pat.font.size = Pt(10)
    run_pat2 = p_pat.add_run('Text based Image Generation and Tourist Place Recommendation System with Style Transfer Algorithm (Jun 2024 to Sep 2024).')
    run_pat2.font.name = 'Arial'
    run_pat2.font.size = Pt(10)

    # Experience & Projects
    add_header(doc, 'EXPERIENCE & PROJECTS', size=12)
    
    add_experience(doc, 'Knowledge Exchange Project | King’s College London and Grocery Insight', 'Data Scientist Researcher, Jan 2026 to Present', [
        'Built a data pipeline to reduce errors in Large Vision Language Models within retail environments.',
        'Created an automated workflow using object detection and text extraction to replace manual tagging rules.',
        'Grouped visual features using unsupervised learning and built a module to generate direct business insights.'
    ])
    
    add_experience(doc, 'Customer Experience Lab | Kookmin University', 'Undergraduate Research Assistant, Jul 2024 to Aug 2024', [
        'Analyzed streaming platform data to understand user viewing habits.',
        'Used content similarity matching to suggest better content and solve user engagement problems.'
    ])
    
    add_experience(doc, 'Alpha Project | Kookmin University', 'Team Lead, Sep 2024 to Dec 2024', [
        'Led data mining and customer experience research to find new product ideas for younger audiences with LG Hellovision.',
        'Built a pet classification model using image deep learning techniques and Vision Transformers.',
        'Processed data and trained classification models to improve targeted marketing.'
    ])
    
    add_experience(doc, 'DataON Research Data Analytics Competition | Ministry of Science and ICT', 'Team Lead, May 2023 to Oct 2023', [
        'Built a river flow prediction model using cloud observation data.',
        'Prepared data and trained a regression model using XGBoost.',
        'Reached the final round as a top 15 team out of 74 entries.'
    ])
    
    add_experience(doc, 'Athletic Data Analytics Competition | Korea Sports Promotion Foundation', 'Team Lead, Oct 2022 to Dec 2022', [
        'Analyzed sports data using clustering algorithms to find ways to increase daily physical activity.',
        'Discovered that changing public perception of exercise is more effective than building new facilities.',
        'Won second place in the data analysis category.'
    ])

    # Save
    doc.save('Junyong_Park_Resume.docx')
    print('Resume generated successfully.')

if __name__ == '__main__':
    main()
